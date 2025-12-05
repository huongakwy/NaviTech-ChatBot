#!/usr/bin/env python3
"""
File Parser - Parse uploaded files to product data or text documents
Support: 
- Product data: Excel (.xlsx, .xls), CSV, JSON
- Text documents: Word (.docx), PDF, TXT for chatbot knowledge base
"""

import pandas as pd
import json
from typing import List, Dict, Union
import os
import re


class TextDocumentParser:
    """Parse text documents for chatbot knowledge base"""
    
    @staticmethod
    def parse_text(file_path: str) -> Dict:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            filename = os.path.basename(file_path)
            content_clean = content.strip()
            
            return {
                'content': content_clean,
                'filename': filename,
                'type': 'text',
                'chunks': TextDocumentParser.chunk_text(content_clean)
            }
        except Exception as e:
            raise ValueError(f"Error parsing text: {str(e)}")
    
    @staticmethod
    def parse_word_text(file_path: str) -> Dict:
        """Parse Word document as text (not table)"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            content = '\n\n'.join(paragraphs)
            filename = os.path.basename(file_path)
            
            return {
                'content': content,
                'filename': filename,
                'type': 'word_text',
                'chunks': TextDocumentParser.chunk_text(content)
            }
        except ImportError:
            raise ValueError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Error parsing Word text: {str(e)}")
    
    @staticmethod
    def parse_pdf_text(file_path: str) -> Dict:
        """Parse PDF as text (not table)"""
        try:
            import pdfplumber
            
            all_text = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        all_text.append(text.strip())
            
            content = '\n\n'.join(all_text)
            filename = os.path.basename(file_path)
            
            return {
                'content': content,
                'filename': filename,
                'type': 'pdf_text',
                'chunks': TextDocumentParser.chunk_text(content)
            }
        except ImportError:
            raise ValueError("pdfplumber not installed. Run: pip install pdfplumber")
        except Exception as e:
            raise ValueError(f"Error parsing PDF text: {str(e)}")
        except ImportError:
            raise ValueError("pdfplumber not installed. Run: pip install pdfplumber")
        except Exception as e:
            raise ValueError(f"Error parsing PDF text: {str(e)}")
    
    @staticmethod
    def chunk_text(content: str, chunk_token_size: int = 2000, **kwargs) -> List[Dict]:
        """
        Split text into chunks using AutoGen's chunking logic
        
        Based on: https://microsoft.github.io/autogen/0.2/docs/notebooks/agentchat_RetrieveChat_qdrant/
        Uses AutoGen's text splitting utilities for consistent chunking
        
        Args:
            content: Full text content
            chunk_token_size: Max tokens per chunk (default 2000)
            **kwargs: Additional args for compatibility
        
        Returns:
            List of text chunks with metadata
        """
        try:
            from autogen.retrieve_utils import split_text_to_chunks
            
            # AutoGen's split_text_to_chunks returns list of strings
            text_chunks = split_text_to_chunks(content, max_tokens=chunk_token_size)
            
            # Convert to our format
            chunks = []
            for idx, chunk_text in enumerate(text_chunks, 1):
                chunks.append({
                    'chunk_id': idx,
                    'text': chunk_text.strip(),
                    'char_count': len(chunk_text)
                })
            
            return chunks
            
        except ImportError:
            # Fallback to sentence-based chunking if AutoGen not available
            import re
            print("⚠️  AutoGen not available, using fallback chunking")
            
            # Estimate chars from tokens (rough: 1 token ≈ 4 chars for English, 2-3 for Vietnamese)
            chunk_size = chunk_token_size * 3  # Conservative estimate for Vietnamese
            overlap = chunk_token_size // 5 * 3  # 20% overlap
            
            sentences = re.split(r'(?<=[.!?])\s+', content)
            
            chunks = []
            current_chunk = ""
            chunk_num = 0
            
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                    chunk_num += 1
                    chunks.append({
                        'chunk_id': chunk_num,
                        'text': current_chunk.strip(),
                        'char_count': len(current_chunk)
                    })
                    
                    if overlap > 0 and len(current_chunk) >= overlap:
                        current_chunk = current_chunk[-overlap:] + " " + sentence
                    else:
                        current_chunk = sentence
                else:
                    if current_chunk:
                        current_chunk += " " + sentence
                    else:
                        current_chunk = sentence
            
            if current_chunk:
                chunk_num += 1
                chunks.append({
                    'chunk_id': chunk_num,
                    'text': current_chunk.strip(),
                    'char_count': len(current_chunk)
                })
            
            return chunks


class FileParser:
    """Parse various file formats to product structure"""
    
    # Mapping từ column names thường gặp → product fields
    COLUMN_MAPPING = {
        # Title variants
        'title': 'title',
        'name': 'title',
        'product_name': 'title',
        'tên sản phẩm': 'title',
        'tên': 'title',
        'ten san pham': 'title',
        'ten': 'title',
        
        # URL variants
        'url': 'url',
        'link': 'url',
        'product_url': 'url',
        
        # Price variants
        'price': 'price',
        'giá': 'price',
        'gia': 'price',
        'giá bán': 'price',
        'gia ban': 'price',
        
        # Original price
        'original_price': 'original_price',
        'giá gốc': 'original_price',
        'gia goc': 'original_price',
        'giá cũ': 'original_price',
        
        # SKU
        'sku': 'sku',
        'mã sản phẩm': 'sku',
        'ma san pham': 'sku',
        'mã': 'sku',
        'ma': 'sku',
        
        # Brand
        'brand': 'brand',
        'thương hiệu': 'brand',
        'thuong hieu': 'brand',
        'nhãn hiệu': 'brand',
        
        # Category
        'category': 'category',
        'danh mục': 'category',
        'danh muc': 'category',
        'loại': 'category',
        'loai': 'category',
        
        # Description
        'description': 'description',
        'mô tả': 'description',
        'mo ta': 'description',
        'thông tin': 'description',
        'thong tin': 'description',
        
        # Images
        'images': 'images',
        'image': 'images',
        'hình ảnh': 'images',
        'hinh anh': 'images',
        'ảnh': 'images',
        'anh': 'images',
        
        # Currency
        'currency': 'currency',
        'đơn vị': 'currency',
        'don vi': 'currency',
    }
    
    @staticmethod
    def parse_excel(file_path: str) -> List[Dict]:
        """Parse Excel file to products"""
        try:
            df = pd.read_excel(file_path)
            return FileParser._dataframe_to_products(df)
        except Exception as e:
            raise ValueError(f"Error parsing Excel: {str(e)}")
    
    @staticmethod
    def parse_csv(file_path: str) -> List[Dict]:
        """Parse CSV file to products"""
        try:
            # Try different encodings
            encodings = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Cannot decode CSV file with any known encoding")
            
            return FileParser._dataframe_to_products(df)
        except Exception as e:
            raise ValueError(f"Error parsing CSV: {str(e)}")
    
    @staticmethod
    def parse_json(file_path: str) -> List[Dict]:
        """Parse JSON file to products"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Nếu JSON là array of products
            if isinstance(data, list):
                return [FileParser._normalize_product(p) for p in data]
            
            # Nếu JSON có key 'products'
            if isinstance(data, dict) and 'products' in data:
                return [FileParser._normalize_product(p) for p in data['products']]
            
            # Nếu JSON là single product
            return [FileParser._normalize_product(data)]
        
        except Exception as e:
            raise ValueError(f"Error parsing JSON: {str(e)}")
    
    @staticmethod
    def parse_word(file_path: str, use_ai_fallback: bool = True) -> List[Dict]:
        """Parse Word file (.docx) containing tables to products"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            all_products = []
            
            # Parse all tables in document
            for table_idx, table in enumerate(doc.tables):
                try:
                    # Extract table to list of lists
                    data = []
                    keys = None
                    
                    for i, row in enumerate(table.rows):
                        text = [cell.text.strip() for cell in row.cells]
                        
                        if i == 0:
                            # First row as header
                            keys = text
                        else:
                            # Data rows
                            if keys and len(text) == len(keys):
                                data.append(dict(zip(keys, text)))
                    
                    if data:
                        df = pd.DataFrame(data)
                        products = FileParser._dataframe_to_products(df)
                        all_products.extend(products)
                        print(f"  ✓ Parsed table {table_idx + 1}: {len(products)} products")
                
                except Exception as e:
                    print(f"  ⚠️ Skipping table {table_idx + 1}: {str(e)[:50]}")
                    continue
            
            # If no tables found or empty, try AI extraction
            if not all_products and use_ai_fallback:
                print("  ℹ️  No tables found, trying AI extraction...")
                try:
                    from utils.ai_extractor import AIProductExtractor
                    all_products = AIProductExtractor.extract_from_file(file_path)
                    print(f"  ✅ AI extracted {len(all_products)} products")
                except Exception as e:
                    print(f"  ⚠️  AI extraction failed: {str(e)[:100]}")
            
            if not all_products:
                raise ValueError("No products found in Word document (tried both table and AI extraction)")
            
            return all_products
        
        except ImportError:
            raise ValueError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Error parsing Word: {str(e)}")
    
    @staticmethod
    def parse_pdf(file_path: str, use_ai_fallback: bool = True) -> List[Dict]:
        """Parse PDF file containing tables to products"""
        try:
            import pdfplumber
            
            all_products = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        # Extract all tables from page
                        tables = page.extract_tables()
                        
                        for table_num, table in enumerate(tables, 1):
                            if not table or len(table) < 2:
                                continue
                            
                            # First row as header
                            headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(table[0])]
                            
                            # Data rows
                            data = []
                            for row in table[1:]:
                                if row:
                                    row_data = [str(cell).strip() if cell else '' for cell in row]
                                    if len(row_data) == len(headers):
                                        data.append(dict(zip(headers, row_data)))
                            
                            if data:
                                df = pd.DataFrame(data)
                                products = FileParser._dataframe_to_products(df)
                                all_products.extend(products)
                                print(f"  ✓ Parsed page {page_num}, table {table_num}: {len(products)} products")
                    
                    except Exception as e:
                        print(f"  ⚠️ Error on page {page_num}: {str(e)[:50]}")
                        continue
            
            # If no tables found or empty, try AI extraction
            if not all_products and use_ai_fallback:
                print("  ℹ️  No tables found, trying AI extraction...")
                try:
                    from utils.ai_extractor import AIProductExtractor
                    all_products = AIProductExtractor.extract_from_file(file_path)
                    print(f"  ✅ AI extracted {len(all_products)} products")
                except Exception as e:
                    print(f"  ⚠️  AI extraction failed: {str(e)[:100]}")
            
            if not all_products:
                raise ValueError("No products found in PDF (tried both table and AI extraction)")
            
            return all_products
        
        except ImportError:
            raise ValueError("pdfplumber not installed. Run: pip install pdfplumber")
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def _dataframe_to_products(df: pd.DataFrame) -> List[Dict]:
        """Convert pandas DataFrame to product list"""
        products = []
        
        # Normalize column names (lowercase, strip spaces)
        df.columns = df.columns.str.lower().str.strip()
        
        for idx, row in df.iterrows():
            product = {}
            
            # Map columns to product fields
            for col in df.columns:
                field = FileParser.COLUMN_MAPPING.get(col)
                if field:
                    value = row[col]
                    
                    # Handle NaN values
                    if pd.isna(value):
                        continue
                    
                    # Convert to appropriate type
                    if field in ['price', 'original_price']:
                        try:
                            # Remove currency symbols, commas, spaces
                            cleaned = str(value).replace(',', '').replace('.', '').replace('đ', '').replace('₫', '').replace('VND', '').replace(' ', '').strip()
                            # Handle European format (comma as decimal separator)
                            if ',' in str(value):
                                cleaned = cleaned.replace(',', '.')
                            product[field] = float(cleaned) if cleaned else 0
                        except:
                            product[field] = 0
                    
                    elif field == 'images':
                        # Split images by comma or semicolon if string
                        if isinstance(value, str):
                            separators = [',', ';', '|']
                            images = [value]
                            for sep in separators:
                                if sep in value:
                                    images = [img.strip() for img in value.split(sep)]
                                    break
                            product[field] = images
                        else:
                            product[field] = [str(value)]
                    
                    else:
                        product[field] = str(value).strip()
            
            # Set defaults
            if 'currency' not in product:
                product['currency'] = 'VND'
            
            if 'url' not in product:
                # Generate pseudo URL from title if missing
                product['url'] = ''
            
            if 'title' in product and product['title']:  # Only add if has title
                products.append(FileParser._normalize_product(product))
        
        return products
    
    @staticmethod
    def _normalize_product(product: Dict) -> Dict:
        """Normalize product data to standard format"""
        normalized = {
            'url': product.get('url', ''),
            'title': product.get('title', ''),
            'price': product.get('price', 0),
            'original_price': product.get('original_price', 0),
            'currency': product.get('currency', 'VND'),
            'sku': product.get('sku', ''),
            'brand': product.get('brand', ''),
            'category': product.get('category', ''),
            'description': product.get('description', ''),
            'availability': product.get('availability', ''),
            'images': product.get('images', [])
        }
        
        # Ensure images is list
        if isinstance(normalized['images'], str):
            normalized['images'] = [normalized['images']]
        
        # Ensure prices are numeric
        try:
            normalized['price'] = float(normalized['price']) if normalized['price'] else 0
        except:
            normalized['price'] = 0
        
        try:
            normalized['original_price'] = float(normalized['original_price']) if normalized['original_price'] else 0
        except:
            normalized['original_price'] = 0
        
        return normalized
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """
        Detect if file is 'product' (structured data) or 'document' (text content)
        Returns: 'product' or 'document'
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        # CSV, Excel, JSON are primarily product data
        if ext in ['.csv', '.xlsx', '.xls', '.json']:
            # Try to detect if it has product columns
            try:
                if ext in ['.xlsx', '.xls']:
                    df = pd.read_excel(file_path, nrows=0)
                elif ext == '.csv':
                    df = pd.read_csv(file_path, nrows=0)
                elif ext == '.json':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        if isinstance(data, list) and len(data) > 0:
                            df = pd.DataFrame([data[0]])
                        else:
                            return 'document'
                
                # Check if has product-related columns
                columns = df.columns.str.lower().str.strip().tolist()
                product_keywords = ['title', 'name', 'price', 'product', 'sku', 'category', 'brand', 'tên', 'giá', 'sản phẩm']
                
                if any(keyword in ' '.join(columns) for keyword in product_keywords):
                    return 'product'
                else:
                    return 'document'
            except:
                return 'product'  # Default to product for these extensions
        
        # TXT files are documents
        elif ext == '.txt':
            return 'document'
        
        # Word and PDF can be either - need to check content
        elif ext in ['.docx', '.pdf']:
            try:
                if ext == '.docx':
                    # Check if has tables (likely product) or mostly text (likely document)
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        if len(doc.tables) > 0:
                            return 'product'
                        else:
                            return 'document'
                    except:
                        return 'document'
                
                elif ext == '.pdf':
                    # Check if has tables (likely product) or mostly text (likely document)
                    try:
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            if len(pdf.pages) > 0:
                                page = pdf.pages[0]
                                tables = page.extract_tables()
                                if tables and len(tables) > 0:
                                    return 'product'
                                else:
                                    return 'document'
                    except:
                        return 'document'
            except:
                return 'document'
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    @staticmethod
    def parse_file(file_path: str) -> List[Dict]:
        """Auto-detect file type and parse as product data"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in ['.xlsx', '.xls']:
            return FileParser.parse_excel(file_path)
        elif ext == '.csv':
            return FileParser.parse_csv(file_path)
        elif ext == '.json':
            return FileParser.parse_json(file_path)
        elif ext == '.docx':
            return FileParser.parse_word(file_path)
        elif ext == '.pdf':
            return FileParser.parse_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: .xlsx, .xls, .csv, .json, .docx, .pdf")
    
    @staticmethod
    def parse_file_as_text(file_path: str) -> Dict:
        """Parse file as text document (for chatbot knowledge base)"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return TextDocumentParser.parse_text(file_path)
        elif ext == '.docx':
            return TextDocumentParser.parse_word_text(file_path)
        elif ext == '.pdf':
            return TextDocumentParser.parse_pdf_text(file_path)
        elif ext in ['.csv', '.xlsx', '.xls', '.json']:
            # For structured files, try to convert to readable text
            try:
                products = FileParser.parse_file(file_path)
                content = "\n\n".join([
                    f"Sản phẩm: {p.get('title', 'N/A')}\n"
                    f"Giá: {p.get('price', 0)} {p.get('currency', 'VND')}\n"
                    f"Mô tả: {p.get('description', 'N/A')}\n"
                    f"Danh mục: {p.get('category', 'N/A')}"
                    for p in products
                ])
                
                filename = os.path.basename(file_path)
                return {
                    'filename': filename,
                    'content': content,
                    'chunks': TextDocumentParser.chunk_text(content)
                }
            except:
                raise ValueError(f"Cannot parse {ext} file as text document")
        else:
            raise ValueError(f"Unsupported file type for text parsing: {ext}")


if __name__ == '__main__':
    # Test parser
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            products = FileParser.parse_file(file_path)
            print(f"✅ Parsed {len(products)} products from {file_path}")
            if products:
                print("\nFirst product:")
                print(json.dumps(products[0], indent=2, ensure_ascii=False))
        except Exception as e:
            print(f"❌ Error: {e}")
